"""Type detection and text extraction.

Extraction is deliberately per page: every downstream citation carries a page
number, so page identity has to survive from the file all the way to the
evidence panel. A parser that returned one flat string would make the
"click the citation, jump to the page" behaviour impossible to build.

Nothing here calls a model. Deciding that a page needs a vision pass is done
by looking at the page's own structure -- how much text it has, how many
images and vector drawings it carries -- because that decision has to be
cheap enough to make for every page of every upload.
"""

from __future__ import annotations

import contextlib
import csv
import io
import json
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

logger = logging.getLogger("workbench.documents")

DocumentKind = Literal[
    "pdf_text", "pdf_scanned", "image", "spreadsheet", "office_text", "text", "unknown"
]

# Below this many characters a PDF page is treated as having no usable text
# layer. A scanned page usually yields a handful of stray glyphs rather than
# nothing at all, so an exact-zero test would miss most real scans.
NATIVE_TEXT_FLOOR = 80

# A page with this little text but real graphics on it is a drawing or a
# photographed form, not a text page: OCR alone would lose the layout, so it
# is flagged for the vision model instead.
VISION_TEXT_CEILING = 400

# Plain-text uploads have no pages, so they are cut into page-sized units.
# This is only for citation granularity -- chunking happens separately.
SYNTHETIC_PAGE_CHARS = 3000


@dataclass
class ExtractedPage:
    page_number: int
    text: str = ""
    # "native" once text came from a text layer, "pending_ocr" while this
    # module has produced pixels but not yet run OCR over them.
    ocr_status: str = "native"
    needs_vision: bool = False
    # PNG bytes, present only when the page has to be looked at, not read.
    image: bytes | None = field(default=None, repr=False)


@dataclass
class ExtractionResult:
    kind: DocumentKind
    pages: list[ExtractedPage]
    detail: str = ""


PDF_MIMES = {"application/pdf"}
IMAGE_MIMES = {"image/png", "image/jpeg", "image/tiff", "image/bmp", "image/webp"}
SPREADSHEET_MIMES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
}
TEXT_MIMES = {"text/plain", "text/csv", "application/json", "text/markdown"}

_SUFFIX_KINDS: dict[str, DocumentKind] = {
    ".pdf": "pdf_text",
    ".png": "image",
    ".jpg": "image",
    ".jpeg": "image",
    ".tif": "image",
    ".tiff": "image",
    ".xlsx": "spreadsheet",
    ".xls": "spreadsheet",
    ".docx": "office_text",
    ".txt": "text",
    ".csv": "text",
    ".json": "text",
    ".md": "text",
}


def detect_kind(mime_type: str, filename: str = "") -> DocumentKind:
    """Classify by declared type, falling back to the extension.

    A PDF is reported as ``pdf_text`` here; whether it is really a scan is
    only knowable after opening it, and ``extract`` corrects the kind.
    """
    mime = (mime_type or "").lower()
    if mime in PDF_MIMES:
        return "pdf_text"
    if mime in IMAGE_MIMES:
        return "image"
    if mime in SPREADSHEET_MIMES:
        return "spreadsheet"
    if mime in DOCX_MIMES:
        return "office_text"
    if mime in TEXT_MIMES:
        return "text"
    return _SUFFIX_KINDS.get(Path(filename).suffix.lower(), "unknown")


def extract(path: Path, mime_type: str, filename: str = "") -> ExtractionResult:
    """Turn a stored file into pages of text, plus pixels where needed."""
    kind = detect_kind(mime_type, filename or path.name)

    if kind in {"pdf_text", "pdf_scanned"}:
        return _extract_pdf(path)
    if kind == "image":
        return _extract_image(path)
    if kind == "spreadsheet":
        return _extract_spreadsheet(path)
    if kind == "office_text":
        return _extract_docx(path)
    if kind == "text":
        return _extract_plain(path)

    return ExtractionResult(
        kind="unknown",
        pages=[],
        detail=f"no extractor for {mime_type or 'unknown type'}",
    )


# --- PDF ------------------------------------------------------------------


def _extract_pdf(path: Path) -> ExtractionResult:
    import fitz  # PyMuPDF

    pages: list[ExtractedPage] = []
    scanned = 0

    with fitz.open(path) as document:
        for index, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            has_graphics = bool(page.get_images(full=False)) or (
                len(page.get_drawings()) > 20
            )

            if len(text) < NATIVE_TEXT_FLOOR:
                # No usable text layer: hand OCR the pixels rather than
                # storing an empty page and pretending the document is short.
                scanned += 1
                pages.append(
                    ExtractedPage(
                        page_number=index,
                        text=text,
                        ocr_status="pending_ocr",
                        needs_vision=has_graphics,
                        image=_render(page),
                    )
                )
                continue

            # A text layer that is thin next to a lot of graphics is a drawing
            # with a title block: readable, but not understandable from text.
            needs_vision = has_graphics and len(text) < VISION_TEXT_CEILING
            pages.append(
                ExtractedPage(
                    page_number=index,
                    text=text,
                    ocr_status="native",
                    needs_vision=needs_vision,
                    image=_render(page) if needs_vision else None,
                )
            )

    kind: DocumentKind = (
        "pdf_scanned" if pages and scanned >= len(pages) / 2 else "pdf_text"
    )
    return ExtractionResult(
        kind=kind,
        pages=pages,
        detail=f"{len(pages)} page(s), {scanned} without a text layer",
    )


def _render(page, dpi: int = 200) -> bytes | None:
    """Rasterise a page for OCR or a vision model.

    200 DPI is the usual floor for reliable OCR of 10-12pt print; going higher
    costs memory for no accuracy gain on the scans this system sees.
    """
    import fitz

    try:
        pixmap = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
        return pixmap.tobytes("png")
    except Exception as exc:
        logger.warning("could not rasterise page %s: %s", page.number, exc)
        return None


# --- images ---------------------------------------------------------------


def _extract_image(path: Path) -> ExtractionResult:
    # A standalone image is always both OCR'd and flagged for vision: an
    # uploaded photograph of a P&ID is exactly the case the spec says must
    # not be handled by OCR alone.
    page = ExtractedPage(
        page_number=1,
        text="",
        ocr_status="pending_ocr",
        needs_vision=True,
        image=path.read_bytes(),
    )
    return ExtractionResult(kind="image", pages=[page], detail="single image")


# --- spreadsheets ---------------------------------------------------------


def _extract_spreadsheet(path: Path) -> ExtractionResult:
    from openpyxl import load_workbook

    pages: list[ExtractedPage] = []
    # read_only keeps a large inspection log from being held in memory whole;
    # data_only takes the cached result of a formula rather than its source.
    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        for index, sheet in enumerate(workbook.worksheets, start=1):
            lines = [f"# Sheet: {sheet.title}"]
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell) for cell in row if cell is not None]
                if cells:
                    lines.append(" | ".join(cells))
            pages.append(
                ExtractedPage(page_number=index, text="\n".join(lines).strip())
            )
    finally:
        workbook.close()

    return ExtractionResult(
        kind="spreadsheet", pages=pages, detail=f"{len(pages)} sheet(s)"
    )


# --- docx -----------------------------------------------------------------


def _extract_docx(path: Path) -> ExtractionResult:
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    blocks = [para.text for para in document.paragraphs if para.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return ExtractionResult(
        kind="office_text",
        pages=_paginate("\n".join(blocks)),
        detail=f"{len(blocks)} block(s)",
    )


# --- plain text -----------------------------------------------------------


def _extract_plain(path: Path) -> ExtractionResult:
    raw = path.read_text(encoding="utf-8", errors="replace")
    suffix = path.suffix.lower()

    if suffix == ".json":
        # Pretty-printed so the chunker has line boundaries to split on.
        # Malformed JSON is still text worth indexing, so it is kept as-is.
        with contextlib.suppress(json.JSONDecodeError):
            raw = json.dumps(json.loads(raw), indent=2, ensure_ascii=False)
    elif suffix == ".csv":
        rows = list(csv.reader(io.StringIO(raw)))
        raw = "\n".join(" | ".join(cell for cell in row) for row in rows)

    return ExtractionResult(kind="text", pages=_paginate(raw), detail="plain text")


def _paginate(text: str) -> list[ExtractedPage]:
    """Cut formats that have no pages into citable units on line boundaries."""
    text = text.strip()
    if not text:
        return []

    pages: list[ExtractedPage] = []
    buffer: list[str] = []
    length = 0

    for line in text.splitlines():
        buffer.append(line)
        length += len(line) + 1
        if length >= SYNTHETIC_PAGE_CHARS:
            pages.append(
                ExtractedPage(page_number=len(pages) + 1, text="\n".join(buffer))
            )
            buffer, length = [], 0

    if buffer:
        pages.append(ExtractedPage(page_number=len(pages) + 1, text="\n".join(buffer)))
    return pages
