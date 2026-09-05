"""DOCX generation.

Deterministic by design: the model supplies validated JSON and this builds the
document. Nothing here asks a model anything, so the same content always
produces the same file, and a broken document is a bug in ninety lines of
code rather than an unreproducible generation.

The citation line under each finding is the point of the whole deliverable.
An approval note whose findings cannot be traced back to a page is not
reviewable, so the generator writes the reference even when it is unflattering
-- a finding with no citation is labelled as unsupported rather than left to
look like the others.
"""

from __future__ import annotations

from typing import ClassVar

from app.artifacts.content import ApprovalNoteContent, Finding
from app.artifacts.store import artifact_store
from app.tools.base import ToolContext, ToolResult

SEVERITY_ORDER = {"critical": 0, "major": 1, "minor": 2, "informational": 3}


class DocxGenerateTool:
    name = "docx.generate"
    description = (
        "Build an approval note as a Word document from structured content. "
        "Supply title, summary, findings (each with severity and citations) "
        "and recommendations. Do not attempt to write the file yourself."
    )
    risk_level = "low"
    # Finalising a deliverable is where a human should look. The orchestrator
    # raises the gate before this runs, so the tool itself stays declarative.
    requires_approval = False
    input_schema: ClassVar[dict] = {
        "type": "object",
        "properties": {
            "title": {"type": "string"},
            "summary": {"type": "string"},
            "findings": {"type": "array"},
            "recommendations": {"type": "array"},
            "filename": {"type": "string"},
        },
        "required": ["title", "summary"],
    }

    def execute(self, args: dict, context: ToolContext) -> ToolResult:
        try:
            content = ApprovalNoteContent.model_validate(
                {
                    "title": args["title"],
                    "summary": args["summary"],
                    "findings": args.get("findings") or [],
                    "recommendations": args.get("recommendations") or [],
                }
            )
        except Exception as exc:
            return ToolResult.failed(f"Content did not validate: {exc}")

        filename = args.get("filename") or "approval_note.docx"
        if not filename.endswith(".docx"):
            filename = f"{filename}.docx"

        payload = build_approval_note(content)
        record = artifact_store.save(
            task_id=context.task_id,
            artifact_type="docx",
            filename=filename,
            payload=payload,
        )
        return ToolResult(
            ok=True,
            data={
                "artifact_id": str(record.id),
                "filename": filename,
                "size_bytes": record.size_bytes,
            },
            detail=f"generated {filename} ({record.size_bytes} bytes)",
        )


def build_approval_note(content: ApprovalNoteContent) -> bytes:
    """Render the note. Pure: content in, bytes out."""
    import io

    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading(content.title, level=0)

    document.add_heading("Summary", level=1)
    document.add_paragraph(content.summary)

    if content.findings:
        document.add_heading("Findings", level=1)
        ordered = sorted(
            content.findings,
            key=lambda finding: SEVERITY_ORDER.get(finding.severity, 9),
        )
        for index, finding in enumerate(ordered, start=1):
            paragraph = document.add_paragraph()
            label = paragraph.add_run(f"{index}. [{finding.severity.upper()}] ")
            label.bold = True
            paragraph.add_run(finding.statement)
            _write_citations(document, finding, Pt)

    if content.recommendations:
        document.add_heading("Recommendations", level=1)
        for recommendation in content.recommendations:
            document.add_paragraph(recommendation, style="List Number")

    if content.references:
        document.add_heading("References", level=1)
        for reference in content.references:
            section = f", {reference.section}" if reference.section else ""
            document.add_paragraph(
                f"[{reference.document_name}, p.{reference.page}{section}]",
                style="List Bullet",
            )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _write_citations(document, finding: Finding, Pt) -> None:
    """Attribute a finding, or say plainly that nothing supports it."""
    line = document.add_paragraph()
    line.paragraph_format.left_indent = Pt(24)
    run = line.add_run()
    run.italic = True
    run.font.size = Pt(9)

    if not finding.citations:
        run.text = "Source: not supported by a retrieved document."
        return

    run.text = "Source: " + "; ".join(
        f"[{citation.document_name}, p.{citation.page}"
        + (f", {citation.section}" if citation.section else "")
        + "]"
        for citation in finding.citations
    )
